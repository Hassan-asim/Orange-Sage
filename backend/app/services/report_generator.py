"""
Report Generator Service for Orange Sage
Handles report generation in various formats
"""

import logging
import os
import uuid
from datetime import datetime
from typing import Dict, List, Any, Optional
from sqlalchemy.orm import Session

from app.models.scan import Scan
from app.models.finding import Finding, SeverityLevel
from app.models.report import Report, ReportFormat, ReportStatus
from app.core.config import settings

logger = logging.getLogger(__name__)


class ReportGenerator:
    """Service for generating security assessment reports"""
    
    def __init__(self):
        self.reports_dir = settings.REPORTS_DIR
        self._ensure_reports_dir()
    
    def _ensure_reports_dir(self):
        """Ensure reports directory exists"""
        if not os.path.exists(self.reports_dir):
            os.makedirs(self.reports_dir, exist_ok=True)
    
    async def generate_report(
        self,
        db: Session,
        scan_id: int,
        format: ReportFormat,
        options: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """Generate a report for a scan"""
        try:
            # Get scan and findings
            scan = db.query(Scan).filter(Scan.id == scan_id).first()
            if not scan:
                raise ValueError(f"Scan {scan_id} not found")
            
            findings = db.query(Finding).filter(Finding.scan_id == scan_id).all()
            
            # Create report record
            report = Report(
                name=f"Security Assessment Report - {scan.name}",
                format=format,
                status=ReportStatus.PENDING,
                scan_id=scan_id,
                include_charts=options.get("include_charts", True),
                include_pocs=options.get("include_pocs", True),
                branding=options.get("branding", "Orange Sage")
            )
            
            db.add(report)
            db.commit()
            
            # Generate report asynchronously
            import asyncio
            asyncio.create_task(self._generate_report_async(report, scan, findings, options))
            
            return {
                "report_id": report.id,
                "status": "generating",
                "message": "Report generation started"
            }
            
        except Exception as e:
            logger.error(f"Error generating report for scan {scan_id}: {e}")
            raise
    
    async def _generate_report_async(
        self,
        report: Report,
        scan: Scan,
        findings: List[Finding],
        options: Dict[str, Any]
    ):
        """Generate report asynchronously"""
        try:
            # Update report status
            report.status = ReportStatus.GENERATING
            db = report.scan.project.owner  # Get database session
            db.commit()
            
            # Generate report based on format
            if report.format == ReportFormat.PDF:
                file_path = await self._generate_pdf_report(report, scan, findings, options)
            elif report.format == ReportFormat.DOCX:
                file_path = await self._generate_docx_report(report, scan, findings, options)
            elif report.format == ReportFormat.HTML:
                file_path = await self._generate_html_report(report, scan, findings, options)
            else:
                raise ValueError(f"Unsupported report format: {report.format}")
            
            # Update report with file info
            report.status = ReportStatus.COMPLETED
            report.storage_key = file_path
            report.file_size = os.path.getsize(file_path)
            report.generated_at = datetime.utcnow()
            report.download_url = f"/api/v1/reports/{report.id}/download"
            report.expires_at = datetime.utcnow() + timedelta(hours=24)
            
            db.commit()
            
            logger.info(f"Generated report {report.id} successfully")
            
        except Exception as e:
            logger.error(f"Error generating report {report.id}: {e}")
            report.status = ReportStatus.FAILED
            report.generation_error = str(e)
            db.commit()
    
    async def _generate_pdf_report(
        self,
        report: Report,
        scan: Scan,
        findings: List[Finding],
        options: Dict[str, Any]
    ) -> str:
        """Generate PDF report"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            
            # Create PDF file
            file_path = os.path.join(self.reports_dir, f"report_{report.id}.pdf")
            doc = SimpleDocTemplate(file_path, pagesize=letter)
            
            # Get styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=1
            )
            
            # Build content
            content = []
            
            # Title
            content.append(Paragraph("Orange Sage Security Assessment Report", title_style))
            content.append(Spacer(1, 20))
            
            # Executive Summary
            content.append(Paragraph("Executive Summary", styles['Heading2']))
            content.append(Paragraph(f"Target: {scan.target.value if scan.target else 'N/A'}", styles['Normal']))
            content.append(Paragraph(f"Scan Date: {scan.created_at.strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
            content.append(Paragraph(f"Total Findings: {len(findings)}", styles['Normal']))
            content.append(Spacer(1, 20))
            
            # Findings Summary Table
            if findings:
                content.append(Paragraph("Findings Summary", styles['Heading2']))
                
                # Count findings by severity
                severity_counts = {}
                for finding in findings:
                    severity = finding.severity.value
                    severity_counts[severity] = severity_counts.get(severity, 0) + 1
                
                # Create summary table
                summary_data = [["Severity", "Count"]]
                for severity in ["critical", "high", "medium", "low", "info"]:
                    count = severity_counts.get(severity, 0)
                    summary_data.append([severity.title(), str(count)])
                
                summary_table = Table(summary_data)
                summary_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                content.append(summary_table)
                content.append(Spacer(1, 20))
            
            # Detailed Findings
            if findings:
                content.append(Paragraph("Detailed Findings", styles['Heading2']))
                
                for i, finding in enumerate(findings, 1):
                    content.append(Paragraph(f"{i}. {finding.title}", styles['Heading3']))
                    content.append(Paragraph(f"Severity: {finding.severity.value.title()}", styles['Normal']))
                    content.append(Paragraph(f"Type: {finding.vulnerability_type or 'N/A'}", styles['Normal']))
                    content.append(Paragraph(f"Endpoint: {finding.endpoint or 'N/A'}", styles['Normal']))
                    content.append(Paragraph(f"Description: {finding.description}", styles['Normal']))
                    
                    if finding.remediation_text:
                        content.append(Paragraph(f"Remediation: {finding.remediation_text}", styles['Normal']))
                    
                    content.append(Spacer(1, 10))
            
            # Build PDF
            doc.build(content)
            
            return file_path
            
        except Exception as e:
            logger.error(f"Error generating PDF report: {e}")
            raise
    
    async def _generate_docx_report(
        self,
        report: Report,
        scan: Scan,
        findings: List[Finding],
        options: Dict[str, Any]
    ) -> str:
        """Generate DOCX report"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Create document
            doc = Document()
            
            # Title
            title = doc.add_heading('Orange Sage Security Assessment Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(f'Target: {scan.target.value if scan.target else "N/A"}')
            doc.add_paragraph(f'Scan Date: {scan.created_at.strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'Total Findings: {len(findings)}')
            
            # Findings Summary
            if findings:
                doc.add_heading('Findings Summary', level=1)
                
                # Count findings by severity
                severity_counts = {}
                for finding in findings:
                    severity = finding.severity.value
                    severity_counts[severity] = severity_counts.get(severity, 0) + 1
                
                # Create summary table
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Severity'
                hdr_cells[1].text = 'Count'
                
                for severity in ["critical", "high", "medium", "low", "info"]:
                    count = severity_counts.get(severity, 0)
                    row_cells = table.add_row().cells
                    row_cells[0].text = severity.title()
                    row_cells[1].text = str(count)
            
            # Detailed Findings
            if findings:
                doc.add_heading('Detailed Findings', level=1)
                
                for i, finding in enumerate(findings, 1):
                    doc.add_heading(f'{i}. {finding.title}', level=2)
                    doc.add_paragraph(f'Severity: {finding.severity.value.title()}')
                    doc.add_paragraph(f'Type: {finding.vulnerability_type or "N/A"}')
                    doc.add_paragraph(f'Endpoint: {finding.endpoint or "N/A"}')
                    doc.add_paragraph(f'Description: {finding.description}')
                    
                    if finding.remediation_text:
                        doc.add_paragraph(f'Remediation: {finding.remediation_text}')
            
            # Save document
            file_path = os.path.join(self.reports_dir, f"report_{report.id}.docx")
            doc.save(file_path)
            
            return file_path
            
        except Exception as e:
            logger.error(f"Error generating DOCX report: {e}")
            raise
    
    async def _generate_html_report(
        self,
        report: Report,
        scan: Scan,
        findings: List[Finding],
        options: Dict[str, Any]
    ) -> str:
        """Generate HTML report"""
        try:
            # Create HTML content
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Orange Sage Security Assessment Report</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 40px; }}
                    .header {{ text-align: center; margin-bottom: 40px; }}
                    .summary {{ background-color: #f5f5f5; padding: 20px; margin: 20px 0; }}
                    .finding {{ border: 1px solid #ddd; margin: 20px 0; padding: 20px; }}
                    .critical {{ border-left: 5px solid #dc3545; }}
                    .high {{ border-left: 5px solid #fd7e14; }}
                    .medium {{ border-left: 5px solid #ffc107; }}
                    .low {{ border-left: 5px solid #28a745; }}
                    .info {{ border-left: 5px solid #17a2b8; }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1>Orange Sage Security Assessment Report</h1>
                </div>
                
                <div class="summary">
                    <h2>Executive Summary</h2>
                    <p><strong>Target:</strong> {scan.target.value if scan.target else 'N/A'}</p>
                    <p><strong>Scan Date:</strong> {scan.created_at.strftime('%Y-%m-%d %H:%M:%S')}</p>
                    <p><strong>Total Findings:</strong> {len(findings)}</p>
                </div>
                
                <h2>Findings Summary</h2>
                <table border="1" style="border-collapse: collapse; width: 100%;">
                    <tr>
                        <th>Severity</th>
                        <th>Count</th>
                    </tr>
            """
            
            # Count findings by severity
            severity_counts = {}
            for finding in findings:
                severity = finding.severity.value
                severity_counts[severity] = severity_counts.get(severity, 0) + 1
            
            for severity in ["critical", "high", "medium", "low", "info"]:
                count = severity_counts.get(severity, 0)
                html_content += f"<tr><td>{severity.title()}</td><td>{count}</td></tr>"
            
            html_content += """
                </table>
                
                <h2>Detailed Findings</h2>
            """
            
            # Add detailed findings
            for i, finding in enumerate(findings, 1):
                severity_class = finding.severity.value
                html_content += f"""
                <div class="finding {severity_class}">
                    <h3>{i}. {finding.title}</h3>
                    <p><strong>Severity:</strong> {finding.severity.value.title()}</p>
                    <p><strong>Type:</strong> {finding.vulnerability_type or 'N/A'}</p>
                    <p><strong>Endpoint:</strong> {finding.endpoint or 'N/A'}</p>
                    <p><strong>Description:</strong> {finding.description}</p>
                """
                
                if finding.remediation_text:
                    html_content += f"<p><strong>Remediation:</strong> {finding.remediation_text}</p>"
                
                html_content += "</div>"
            
            html_content += """
            </body>
            </html>
            """
            
            # Save HTML file
            file_path = os.path.join(self.reports_dir, f"report_{report.id}.html")
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return file_path
            
        except Exception as e:
            logger.error(f"Error generating HTML report: {e}")
            raise
    
    async def get_report_status(self, report_id: int, db: Session) -> Dict[str, Any]:
        """Get report generation status"""
        report = db.query(Report).filter(Report.id == report_id).first()
        if not report:
            return {"error": "Report not found"}
        
        return {
            "report_id": report_id,
            "status": report.status.value,
            "format": report.format.value,
            "name": report.name,
            "file_size": report.file_size,
            "download_url": report.download_url,
            "expires_at": report.expires_at.isoformat() if report.expires_at else None,
            "generated_at": report.generated_at.isoformat() if report.generated_at else None,
            "error": report.generation_error
        }
    
    async def download_report(self, report_id: int, db: Session) -> Optional[str]:
        """Get report download path"""
        report = db.query(Report).filter(Report.id == report_id).first()
        if not report or report.status != ReportStatus.COMPLETED:
            return None
        
        return report.storage_key
